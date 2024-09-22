# env - all2py310
# pip install python-docx
# pip install windows-tools

import time
from datetime import datetime as dt, timedelta
from platform import node
from struct import unpack
from winreg import OpenKeyEx, QueryValueEx, HKEY_LOCAL_MACHINE, QueryInfoKey, EnumKey, KEY_READ

from docx import Document

from windows_tools import product_key

# Узнаем версию ОС
def winreg_os() -> dict:
    win_info = dict()
    if comp_info := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SYSTEM\CurrentControlSet\Control\ComputerName\ComputerName"):
        win_info.update({'ComputerName': QueryValueEx(comp_info, 'ComputerName')[0]})
    if comp_shutdown := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SYSTEM\CurrentControlSet\Control\Windows"):
        shutdown_time_bin = QueryValueEx(comp_shutdown, 'ShutdownTime')[0]
        shutdown_time = (dt(1601, 1, 1) + timedelta(microseconds=float(unpack("<Q", shutdown_time_bin)[0]) / 10)). \
            strftime('%Y-%m-%d %H:%M:%S')
        win_info.update({'ShutdownTime': shutdown_time})
    if win_ver := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion"):
        for key in ["ProductName", "EditionID", "DisplayVersion", "CurrentBuild", "UBR", "InstallDate",
                    "RegisteredOwner"]:
            try:
                if key == "InstallDate":
                    win_info.update({key: str(dt.fromtimestamp(QueryValueEx(win_ver, f'{key}')[0]))})
                else:
                    win_info.update({key: QueryValueEx(win_ver, f'{key}')[0]})
            except FileNotFoundError:
                continue
    if tz_key := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SYSTEM\CurrentControlSet\Control\TimeZoneInformation"):
        win_info.update({"TimeZone": QueryValueEx(tz_key, 'TimeZoneKeyName')[0]})
    if pkey := product_key.get_windows_product_key_from_reg():
        win_info.update({"ActivateKey": pkey})
    elif pkey := product_key.get_windows_product_key_from_wmi():
        win_info.update({"ActivateKey": pkey})
    else:
        win_info.update({"ActivateKey": "No Key"})
    return win_info if win_info else False

wmic_info = ""

def print_wmic(part, dict_info):
    global wmic_info
    synonyms = {"ComputerName": "Имя компьютера", "Caption": "Название", "InstallDate": "Дата установки",
                "LastBootUpTime": "Время последней загрузки", "Version": "Версия",
                "WindowsDirectory": "Директория Windows", "TimeZone": "Часовой пояс", "UserName": "Имя пользователя",
                "Manufacturer": "Производитель", "Name": "Название", "Product": "Изделие",
                "MaxClockSpeed": "Максимальная тактовая частота", "SocketDesignation": "Название сокета",
                "NumberOfPhysicalProcessors": "Количество физических процессоров", "VideoProcessor": "Видеопроцессор",
                "NumberOfLogicalProcessors": "Количество логических процессоров", "Capacity": "Емкость",
                "AdapterRAM": "Оперативная память адаптера", "CurrentRefreshRate": "Текущая частота обновления",
                "Resolution": "Разрешение", "TotalPhysicalMemory": "Общий объем физической памяти", "Socket": "Сокет",
                "ConfiguredClockSpeed": "Настроенная тактовая частота", "PartNumber": "Номер партии",
                "SerialNumber": "Серийный номер", "DeviceID": "Идентификатор устройства", "MediaType": "Тип носителя",
                "FirmwareRevision": "Ревизия прошивки", "Partitions": "Разделы", "Size": "Объем", "Drive": "Диск",
                "VolumeName": "Имя тома", "VolumeSerialNumber": "Серийный номер тома", "MACAddress": "MAC-адрес",
                "NetConnectionID": "Идентификатор сетевого подключения", "DHCPServer": "DHCP-сервер",
                "IPAddress": "IP-адрес", "BuildNumber": "Номер сборки", "ID": "Идентификатор", "Status": "Статус",
                "DefaultIPGateway": "IP-адрес шлюза по-умолчанию", "DNSHostName": "DNS Имя хоста",
                "IPv4Address": "IPv4-адрес", "IPv6Address": "IPv6-адрес", "IPSubnet": "Маска подсети",
                "ServiceName": "Название службы", "CurrentBuild": "Текущая сборка", "UBR": "Номер версии",
                "RegisteredOwner": "Имя пользователя", "ActivateKey": "Ключ активации",
                "SystemBiosVersion": "Версия Bios системы", "BIOSVendor": "Производитель", "BIOSVersion": "Версия",
                "BIOSReleaseDate": "Дата выпуска релиза", "ShutdownTime": "Время выключения", "ProductName": "Название",
                "EditionID": "Идентификатор редакции", "DisplayVersion": "Версия для отображения",
                "SystemManufacturer": "Производитель", "SystemProductName": "Название сокета",
                "CoreCount": "Количество ядер", "ProcessorNameString": "Название", "Identifier": "Идентификатор",
                "VendorIdentifier": "Производитель", "~MHz": "Тактовая частота", "Vendor": "Производитель",
                "Model": "Модель", "Revision": "Ревизия", "Description": "Название",
                "NetCfgInstanceId": "Идентификатор", "DhcpDefaultGateway": "Шлюз по-умолчанию",
                "DhcpIPAddress": "IP-адрес"}
    part += f'{"-" * 50}\n'
    for key in dict_info:
        if type(dict_info[key]) == dict:
            for item in dict_info[key]:
                part += f'{synonyms[item]}: {dict_info[key][item]}\n'
            part += "\n"
        else:
            part += f'{synonyms[key]}: {dict_info[key]}\n'
    print(part)
    wmic_info += f'{part}\n'

def main():
    global wmic_info
    t = time.monotonic()
    document = Document()
    document.add_heading(f'Сводная информация о компьютере: {node()}')

    if os_info := winreg_os():
        print_wmic("Информация об операционной системе\n", os_info)

    document.add_paragraph(wmic_info)
    document.save(f'{node()}.docx')
    print(f"Собранная информация сохранена в файл: {node()}.docx")
    print(f'\nВремя работы скрипта: {time.monotonic() - t} с.')

if __name__ == "__main__":
    main()